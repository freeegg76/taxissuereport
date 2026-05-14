import re
import os
import platform
from pathlib import Path
from sqlalchemy.orm import Session

from app.db import crud
from app.db.models import Report

_EXPORT_DIR = Path(__file__).parents[3] / "output" / "exports"
_EXPORT_DIR.mkdir(parents=True, exist_ok=True)

_CITATION_RE = re.compile(r'\[([^\]]+)\]')

# Korean font paths
def _get_korean_fonts() -> tuple[str, str]:
    if platform.system() == "Windows":
        regular = r"C:\Windows\Fonts\malgun.ttf"
        bold = r"C:\Windows\Fonts\malgunbd.ttf"
    else:
        # Linux fallback (NotoSansCJK or UnBatang)
        for path in [
            "/usr/share/fonts/truetype/noto/NotoSansCJKkr-Regular.otf",
            "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
            "/usr/share/fonts/truetype/unfonts-core/UnBatang.ttf",
        ]:
            if os.path.exists(path):
                return path, path
        regular = ""
        bold = ""
    return regular, bold


# ── DOCX ─────────────────────────────────────────────────────────────────────

def generate_docx(report: Report, db: Session) -> dict:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # Set default body font to Malgun Gothic for Korean support
    _set_doc_korean_font(doc)

    markdown = report.full_report_markdown or ""
    file_name = f"report_{report.report_id}.docx"
    file_path = str(_EXPORT_DIR / file_name)

    for line in markdown.split("\n"):
        if line.startswith("# "):
            p = doc.add_heading(line[2:], level=1)
            _set_para_korean_font(p)
        elif line.startswith("## "):
            p = doc.add_heading(line[3:], level=2)
            _set_para_korean_font(p)
        elif line.startswith("### "):
            p = doc.add_heading(line[4:], level=3)
            _set_para_korean_font(p)
        elif line.startswith("> "):
            p = doc.add_paragraph(line[2:], style="Quote")
            _set_para_korean_font(p)
        elif line.startswith("※ 검토 의견:"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(12)
            run = p.add_run(line)
            run.font.color.rgb = RGBColor(0x80, 0x60, 0x00)
            run.font.italic = True
            _set_run_korean_font(run)
        elif line.startswith("|") and "|" in line[1:]:
            p = doc.add_paragraph(line)
            _set_para_korean_font(p)
        elif line.strip():
            p = doc.add_paragraph()
            _add_styled_runs(p, line)
        else:
            doc.add_paragraph()

    doc.save(file_path)
    ef = crud.save_exported_file(db, report.report_id, "docx", file_path, file_name)
    return {"file_id": ef.file_id, "file_name": file_name, "file_path": file_path}


def _set_doc_korean_font(doc) -> None:
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    for style_name in ("Normal", "Heading 1", "Heading 2", "Heading 3"):
        try:
            style = doc.styles[style_name]
            style.font.name = "Malgun Gothic"
            # Set East Asian font via XML (required for Korean glyphs)
            rPr = style.element.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rPr.insert(0, rFonts)
            rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
            rFonts.set(qn("w:ascii"), "Malgun Gothic")
            rFonts.set(qn("w:hAnsi"), "Malgun Gothic")
        except Exception:
            pass


def _set_run_korean_font(run) -> None:
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    run.font.name = "Malgun Gothic"
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    rFonts.set(qn("w:ascii"), "Malgun Gothic")
    rFonts.set(qn("w:hAnsi"), "Malgun Gothic")


def _set_para_korean_font(para) -> None:
    for run in para.runs:
        _set_run_korean_font(run)


def _add_styled_runs(paragraph, text: str) -> None:
    from docx.shared import RGBColor

    parts = _CITATION_RE.split(text)
    for i, part in enumerate(parts):
        if i % 2 == 1:
            run = paragraph.add_run(f"[{part}]")
            run.font.color.rgb = RGBColor(0x00, 0x56, 0xA8)
            run.underline = True
        else:
            run = paragraph.add_run(part)
        _set_run_korean_font(run)


# ── PDF (reportlab) ───────────────────────────────────────────────────────────

def generate_pdf(report: Report, db: Session) -> dict:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.lib.enums import TA_LEFT, TA_CENTER
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, HRFlowable, Table, TableStyle
    )
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.lib import colors

    regular_path, bold_path = _get_korean_fonts()
    if not regular_path or not os.path.exists(regular_path):
        raise RuntimeError(
            "한국어 폰트를 찾을 수 없습니다. "
            "Windows: C:\\Windows\\Fonts\\malgun.ttf, "
            "Linux: NotoSansCJK 또는 UnBatang 설치 필요"
        )

    pdfmetrics.registerFont(TTFont("Korean", regular_path))
    if bold_path and os.path.exists(bold_path):
        pdfmetrics.registerFont(TTFont("Korean-Bold", bold_path))
    else:
        pdfmetrics.registerFont(TTFont("Korean-Bold", regular_path))

    file_name = f"report_{report.report_id}.pdf"
    file_path = str(_EXPORT_DIR / file_name)

    doc = SimpleDocTemplate(
        file_path,
        pagesize=A4,
        leftMargin=20 * mm,
        rightMargin=20 * mm,
        topMargin=20 * mm,
        bottomMargin=20 * mm,
    )

    styles = {
        "h1": ParagraphStyle("h1", fontName="Korean-Bold", fontSize=16, spaceAfter=6,
                             textColor=colors.HexColor("#2c3e50")),
        "h2": ParagraphStyle("h2", fontName="Korean-Bold", fontSize=13, spaceAfter=4,
                             spaceBefore=10, textColor=colors.HexColor("#2c3e50")),
        "h3": ParagraphStyle("h3", fontName="Korean-Bold", fontSize=11, spaceAfter=3,
                             spaceBefore=6, textColor=colors.HexColor("#34495e")),
        "body": ParagraphStyle("body", fontName="Korean", fontSize=10, spaceAfter=4,
                               leading=16),
        "opinion": ParagraphStyle("opinion", fontName="Korean", fontSize=10,
                                  spaceAfter=4, leading=16,
                                  leftIndent=8, borderPadding=(4, 8, 4, 8),
                                  backColor=colors.HexColor("#fef9f0"),
                                  borderColor=colors.HexColor("#f0ad4e"),
                                  borderWidth=0, borderRadius=0),
        "quote": ParagraphStyle("quote", fontName="Korean", fontSize=10,
                                leftIndent=12, spaceAfter=4, leading=16,
                                textColor=colors.HexColor("#555555")),
    }

    story = []
    markdown = report.full_report_markdown or ""

    for line in markdown.split("\n"):
        line_stripped = line.strip()
        if not line_stripped:
            story.append(Spacer(1, 4))
            continue

        # Escape special reportlab characters
        safe = _escape_rl(line_stripped)

        if line.startswith("# "):
            story.append(Paragraph(_escape_rl(line[2:].strip()), styles["h1"]))
        elif line.startswith("## "):
            story.append(Spacer(1, 4))
            story.append(Paragraph(_escape_rl(line[3:].strip()), styles["h2"]))
        elif line.startswith("### "):
            story.append(Paragraph(_escape_rl(line[4:].strip()), styles["h3"]))
        elif line.startswith("> "):
            story.append(Paragraph(_escape_rl(line[2:].strip()), styles["quote"]))
        elif line_stripped.startswith("※ 검토 의견:"):
            story.append(Paragraph(safe, styles["opinion"]))
        elif line_stripped.startswith("---"):
            story.append(HRFlowable(width="100%", color=colors.HexColor("#cccccc")))
        elif line_stripped.startswith("|"):
            # Simple table row — render as plain text
            story.append(Paragraph(safe, styles["body"]))
        else:
            story.append(Paragraph(safe, styles["body"]))

    doc.build(story)
    ef = crud.save_exported_file(db, report.report_id, "pdf", file_path, file_name)
    return {"file_id": ef.file_id, "file_name": file_name, "file_path": file_path}


def _escape_rl(text: str) -> str:
    """Escape special ReportLab XML characters."""
    return (text
            .replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;"))
